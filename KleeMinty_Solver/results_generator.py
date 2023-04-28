import pandas as pd
from docx import Document
from docx.shared import Inches
import os
import json

def add_table(document, title, dataframe):
    # Adiciona Título
    title = document.add_paragraph(title)
    title.alignment = 1

    # Adiciona o dataframe1 na tabela
    table = document.add_table(dataframe.shape[0]+1, dataframe.shape[1])
    table.style = 'Table Grid'

    # Adiciona os cabeçalhos
    for j in range(dataframe.shape[-1]):
        table.cell(0,j).text = dataframe.columns[j]

    # Adiciona os valores
    for i in range(dataframe.shape[0]):
        for j in range(dataframe.shape[-1]):
            table.cell(i+1,j).text = str(dataframe.values[i,j])

    # Centraliza os dados da tabela
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = 1  # Define alinhamento centralizado horizontalmente
            cell.vertical_alignment = 1 # Define alinhamento centralizado horizontalmente

    space = document.add_paragraph('')
    space.add_run().add_break()

def gen_results():
    # Abrindo o arquivo JSON e lendo seu conteúdo
    with open(os.path.dirname(__file__) + '/../dados.json', 'r') as f:
        dados = json.load(f)

    dataframe1 = pd.DataFrame(dados['dataframe1'])

    dataframe2 = pd.DataFrame(dados['dataframe2'])

    dataframe3 = pd.DataFrame(dados['dataframe3'])

    dataframe4 = pd.DataFrame(dados['dataframe4'])

    dataframe5 = pd.DataFrame(dados['dataframe5'])

    dataframe6 = pd.DataFrame(dados['dataframe6'])

    dataframe7 = pd.DataFrame(dados['dataframe7'])

    dataframe8 = pd.DataFrame(dados['dataframe8'])

    dataframe9 = pd.DataFrame(dados['dataframe9'])

    dataframe10 = pd.DataFrame(dados['dataframe10'])

    dataframe11 = pd.DataFrame(dados['dataframe11'])

    dataframe12 = pd.DataFrame(dados['dataframe12'])

    dataframe13 = pd.DataFrame(dados['dataframe13'])

    dataframe14 = pd.DataFrame(dados['dataframe14'])

    dataframe15 = pd.DataFrame(dados['dataframe15'])

    dataframe16 = pd.DataFrame(dados['dataframe16'])



    document = Document()

    #Cria a tabela 1
    title_table_1 = 'Table 1 – Execution Time × Problem Dimension'
    add_table(document=document,title=title_table_1, dataframe=dataframe1)

    #Cria a tabela 2
    title_table_2 = 'Table 2 – Number of iterations × Problem Dimension'
    add_table(document=document,title=title_table_2, dataframe=dataframe2)

    #Cria a tabela 3
    title_table_3 = 'Table 3 – Objective Function Error × Problem Dimension'
    add_table(document=document,title=title_table_3, dataframe=dataframe3)

    #Cria a tabela 4
    title_table_4 = 'Table 4 – Solution Error × Problem Dimension'
    add_table(document=document,title=title_table_4, dataframe=dataframe4)

    #Cria a tabela 5
    title_table_5 = 'Table 5 – Step (Alpha) x Execution Time - Interior Point Algorithm'
    add_table(document=document,title=title_table_5, dataframe=dataframe5)

    #Cria a tabela 6
    title_table_6 = 'Table 6 – Step (Alpha) x Iterations - Interior Point Algorithm'
    add_table(document=document,title=title_table_6, dataframe=dataframe6)

    #Cria a tabela 7
    title_table_7 = 'Table 7 – Step (Alpha) x Objective Function Error (%) - Interior Point Algorithm'
    add_table(document=document,title=title_table_7, dataframe=dataframe7)

    #Cria a tabela 8
    title_table_8 = 'Table 8 – Step (Alpha) x Solution Error (%) - Interior Point Algorithm'
    add_table(document=document,title=title_table_8, dataframe=dataframe8)

    #Cria a tabela 9
    title_table_9 = 'Table 9 – Step (Alpha) x Execution Time - Hybrid Algorithm'
    add_table(document=document,title=title_table_9, dataframe=dataframe9)

    #Cria a tabela 10
    title_table_10 = 'Table 10 – Step (Alpha) x Iterations - Hybrid Algorithm'
    add_table(document=document,title=title_table_10, dataframe=dataframe10)

    #Cria a tabela 11
    title_table_11 = 'Table 11 – Gap x Execution Time - Interior Point Algorithm'
    add_table(document=document,title=title_table_11, dataframe=dataframe11)

    #Cria a tabela 12
    title_table_12 = 'Table 12 – Gap x Iterations - Interior Point Algorithm'
    add_table(document=document,title=title_table_12, dataframe=dataframe12)

    #Cria a tabela 13
    title_table_13 = 'Table 13 – Gap x Objective Function Error (%) - Interior Point Algorithm'
    add_table(document=document,title=title_table_13, dataframe=dataframe13)

    #Cria a tabela 14
    title_table_14 = 'Table 14 – Gap x Solution Error (%) - Interior Point Algorithm'
    add_table(document=document,title=title_table_14, dataframe=dataframe14)

    #Cria a tabela 15
    title_table_15 = 'Table 15 – Gap x Execution Time - Hybrid Algorithm'
    add_table(document=document,title=title_table_15, dataframe=dataframe15)

    #Cria a tabela 16
    title_table_16 = 'Table 16 – Gap x Iterations - Hybrid Algorithm'
    add_table(document=document,title=title_table_16, dataframe=dataframe16)

    # Salva o arquivo em .docx
    document.save(os.path.dirname(__file__) + '/../Results.docx')

    #Abre o documento
    os.startfile(os.path.dirname(__file__) + '/../Results.docx')

    return